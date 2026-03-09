"""
Agreement text extraction from various file formats
Supports PDF, DOCX, and TXT files
"""

import os
from typing import Optional, Tuple
import pdfplumber
from docx import Document


class AgreementParser:
    """Extract text from agreement documents"""

    @staticmethod
    def extract_from_pdf(file_path: str) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text if text.strip() else None
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_docx(file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text if text.strip() else None
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_txt(file_path: str) -> Optional[str]:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text if text.strip() else None
        except Exception as e:
            print(f"Error extracting text from TXT {file_path}: {e}")
            return None

    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """
        Extract text from any supported format
        
        Args:
            file_path: Path to agreement file
            
        Returns:
            Extracted text or None if failed
        """
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return AgreementParser.extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return AgreementParser.extract_from_docx(file_path)
        elif file_ext == '.txt':
            return AgreementParser.extract_from_txt(file_path)
        else:
            print(f"Unsupported file format: {file_ext}")
            return None

    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, str]:
        """
        Validate if file is a supported format
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (is_valid, message)
        """
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        supported = ['.pdf', '.docx', '.txt']
        
        if file_ext not in supported:
            return False, f"Unsupported format. Supported: {', '.join(supported)}"
        
        return True, "File format is supported"
